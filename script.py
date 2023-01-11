import os
from docx import Document

def replace_text_in_docx(file_path, old_text, new_text):
    document = Document(file_path)
    for p in document.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
                for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = para.text.replace(old_text, new_text)
    document.save(file_path)
    
def search_and_replace(root_dir, old_text, new_text):
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.endswith('.docx'):
                file_path = os.path.join(dirpath, filename)
                replace_text_in_docx(file_path, old_text, new_text)

root_dir =  r'c://path'
old_text = 'word1'
new_text = 'word2'
search_and_replace(root_dir, old_text, new_text)

print("All done!")
